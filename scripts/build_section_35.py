from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = "/Users/worakanp/Desktop/go-with-us-1/outputs/บทที่_3_หัวข้อ_3.5_การคำนวณ_Cosine_Similarity.docx"
FONT = "Tahoma"
BLUE = "244A64"
LIGHT = "EAF1F5"
GRAY = "F3F4F6"

def font(run, size=16, bold=False, color="000000"):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    font(r, 18 if level == 1 else 17, True, BLUE)
    return p

def body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.first_line_indent = Cm(1.25)
    font(p.add_run(text), 16)
    return p

def formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(text), 16, True, BLUE)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(text), 16)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, BLUE); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(h), 15, True, "FFFFFF"); set_cell_margins(c)
        if widths: c.width = Cm(widths[i])
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]; c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri % 2: shade(c, GRAY)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(headers) == 2 else WD_ALIGN_PARAGRAPH.LEFT
            font(p.add_run(str(val)), 15); set_cell_margins(c)
            if widths: c.width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(2.54), Cm(2.54)
sec.left_margin, sec.right_margin = Cm(3.0), Cm(2.5)
sec.header_distance, sec.footer_distance = Cm(1.25), Cm(1.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(16)
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.paragraph_format.line_spacing = 1.15

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Go With Us — การออกแบบระบบแนะนำทริป"), 11, color="666666")

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
font(title.add_run("3.5 การออกแบบโครงสร้างการวิเคราะห์และแนะนำข้อมูล"), 22, True, BLUE)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(14)
font(sub.add_run("ระบบจับคู่ผู้ใช้งานกับทริปด้วย Cosine Similarity"), 16, False, "555555")

body(doc, "ระบบ Go With Us ออกแบบกลไกแนะนำทริปโดยใช้การวิเคราะห์ความคล้ายคลึงเชิงโคไซน์ (Cosine Similarity) เพื่อประเมินว่าความต้องการของผู้ใช้งานสอดคล้องกับคุณลักษณะของแต่ละทริปมากน้อยเพียงใด ข้อมูลคำตอบจากแบบสอบถามจะถูกแปลงเป็นเวกเตอร์ความสนใจของผู้ใช้ (User Preference Vector) ส่วนข้อมูลของทริปจะถูกแปลงเป็นเวกเตอร์คุณลักษณะของทริป (Trip Feature Vector) จากนั้นระบบจึงเปรียบเทียบทิศทางของเวกเตอร์ทั้งสองและแสดงผลเป็น Match Score ร้อยละ 0–100")
body(doc, "ระบบพิจารณาคุณลักษณะ 4 ด้าน โดยให้น้ำหนักหมวดหมู่การท่องเที่ยวมากที่สุด เพราะสะท้อนจุดประสงค์หลักของการเดินทาง รองลงมาคืองบประมาณ ระดับกิจกรรม และช่วงเวลาที่ชอบทำกิจกรรม น้ำหนักดังกล่าวเป็นค่าที่ออกแบบจากเหตุผลเชิงโดเมนของระบบ มิใช่ค่าที่ได้รับการยืนยันว่าเหมาะสมที่สุดทางสถิติ และสามารถปรับปรุงได้เมื่อมีข้อมูลพฤติกรรมผู้ใช้จริงมากเพียงพอ")
table(doc, ["คุณลักษณะ", "น้ำหนัก", "เหตุผล"], [
    ("หมวดหมู่การท่องเที่ยว", "35%", "สะท้อนความสนใจหลัก เช่น ทะเล ภูเขา หรือคาเฟ่"),
    ("งบประมาณ", "30%", "ตรวจสอบว่าค่าใช้จ่ายของทริปอยู่ใกล้ระดับที่ผู้ใช้ยอมรับได้"),
    ("ระดับกิจกรรม", "20%", "เปรียบเทียบความหนาแน่นของแผน เช่น เที่ยวสบายหรือกิจกรรมแน่น"),
    ("ช่วงเวลาทำกิจกรรม", "15%", "เปรียบเทียบความชอบช่วงเช้า กลางวัน เย็น และกลางคืน"),
], [4.8, 2.2, 8.5])

heading(doc, "3.5.1 กระบวนการสกัดคุณลักษณะและแปลงข้อมูลเป็นเวกเตอร์")
body(doc, "ข้อมูลข้อความและตัวเลือกจากผู้ใช้ยังไม่สามารถนำเข้าสูตร Cosine Similarity ได้โดยตรง ระบบจึงต้องสกัดคุณลักษณะ (Feature Extraction) และแปลงให้อยู่ในรูปชุดตัวเลข (Vectorization) โดยตำแหน่งเดียวกันของเวกเตอร์ผู้ใช้และเวกเตอร์ทริปต้องมีความหมายเหมือนกัน เช่น ตำแหน่งที่แทนหมวดหมู่ทะเลต้องตรงกันทั้งสองเวกเตอร์")

heading(doc, "3.5.1.1 คุณลักษณะด้านงบประมาณ", 2)
body(doc, "ระบบเปรียบเทียบงบประมาณที่ผู้ใช้ระบุกับงบประมาณของทริปโดยตรง ไม่ได้หารด้วยจำนวนวันเพื่อทำเป็นงบเฉลี่ยต่อวัน และคำว่า ‘ต่อคน’ หรือ ‘ต่อทริป’ ใช้เพื่อแสดงความหมายแก่ผู้ใช้เท่านั้น ไม่ได้เป็นคุณลักษณะอีกมิติหนึ่งในการจับคู่ เนื่องจากงบมีช่วงค่ากว้างตั้งแต่หลักร้อยถึงหลักหมื่น ระบบจึงใช้ลอการิทึมลดอิทธิพลของค่าจำนวนเงินขนาดใหญ่")
formula(doc, "position_budget = [ln(B) − ln(100)] ÷ [ln(50,000) − ln(100)]")
body(doc, "เมื่อ B คืองบประมาณ ระบบจะจำกัดค่าที่ใช้แปลงให้อยู่ในช่วง 100–50,000 บาท ค่า position อยู่ระหว่าง 0–1 และเป็นเพียงค่าระหว่างทาง ไม่ได้อยู่ใน Dot Product โดยตรง จากนั้นระบบเปลี่ยน position เป็นมุมบนส่วนหนึ่งในสี่ของวงกลม และสร้างเวกเตอร์สองมิติ")
formula(doc, "θ_budget = position_budget × π/2")
formula(doc, "Budget Vector = [cos(θ_budget), sin(θ_budget)]")
body(doc, "วิธีนี้ทำให้งบที่ใกล้กันมีทิศทางของเวกเตอร์ใกล้กัน และคะแนนลดลงอย่างต่อเนื่องเมื่องบแตกต่างกัน โดยไม่ทำให้จำนวนเงินครอบงำคุณลักษณะด้านอื่น")

heading(doc, "3.5.1.2 คุณลักษณะด้านระดับกิจกรรม", 2)
body(doc, "ระดับกิจกรรมหมายถึงความหนาแน่นของสถานที่หรือกิจกรรมที่ต้องการทำต่อวัน แบบสอบถามมี 3 ตัวเลือก แต่แปลงเป็นค่า 2, 5 และ 8 เพื่อแทนระดับต่ำ กลาง และสูง ค่าดังกล่าวเป็นตัวแทนระดับ มิได้บังคับว่าทริปต้องมีจำนวนกิจกรรมเท่ากับตัวเลขนั้นทุกครั้ง")
table(doc, ["คำตอบ", "ค่าที่ใช้", "ความหมาย"], [
    ("1–2 สถานที่/กิจกรรมต่อวัน", "2", "เที่ยวสบายและมีเวลาพัก"),
    ("3–4 สถานที่/กิจกรรมต่อวัน", "5", "ความหนาแน่นระดับปานกลาง"),
    ("5 สถานที่/กิจกรรมขึ้นไปต่อวัน", "8", "เน้นกิจกรรมค่อนข้างแน่น"),
], [6.4, 2.5, 6.6])
formula(doc, "position_activity = (A − 1) ÷ 9")
formula(doc, "Activity Vector = [cos(position_activity × π/2), sin(position_activity × π/2)]")
body(doc, "เมื่อผู้ใช้กับทริปมีระดับกิจกรรมเท่ากัน เวกเตอร์จะชี้ไปในทิศทางเดียวกันและได้คะแนนส่วนนี้ใกล้ 1 หากระดับแตกต่างกัน มุมระหว่างเวกเตอร์จะเพิ่มขึ้นและคะแนนจะลดลง")

heading(doc, "3.5.1.3 คุณลักษณะด้านช่วงเวลา", 2)
body(doc, "ผู้ใช้เลือกได้หลายช่วงเวลา ได้แก่ เช้า กลางวัน เย็น และกลางคืน แต่ละตำแหน่งถูกแทนด้วยค่า 1 เมื่อเลือก และ 0 เมื่อไม่เลือก เช่น ผู้ใช้เลือกเช้า เย็น และกลางคืน จะได้เวกเตอร์เริ่มต้น [1, 0, 1, 1] ระบบจะปรับขนาดเวกเตอร์ให้เท่ากับ 1 เพื่อไม่ให้ผู้ที่เลือกหลายตัวเลือกได้เปรียบผู้ที่เลือกน้อย")
formula(doc, "||T|| = √(1² + 0² + 1² + 1²) = √3")
formula(doc, "T_normalized = [1/√3, 0, 1/√3, 1/√3] = [0.5774, 0, 0.5774, 0.5774]")
body(doc, "ดังนั้น 0.5774 มาจาก 1 หารด้วยรากที่สองของ 3 เนื่องจากตัวอย่างนี้เลือก 3 ช่วงเวลา มิใช่ค่าคงที่ที่ระบบตั้งขึ้นเอง")

heading(doc, "3.5.1.4 คุณลักษณะด้านหมวดหมู่การท่องเที่ยว", 2)
body(doc, "ระบบรองรับ 13 หมวด ได้แก่ ทะเล ภูเขา แคมป์ปิ้ง เที่ยวเมือง คาเฟ่ อาหาร แฮงเอาต์ ถ่ายรูป ช้อปปิ้ง คอนเสิร์ต ผจญภัย ไหว้พระ และอื่น ๆ แต่ละหมวดมีตำแหน่งคงที่ในเวกเตอร์ หากเลือกจะมีค่า 1 และหากไม่เลือกจะมีค่า 0 เมื่อเลือกหลายหมวด ระบบจะปรับขนาดเวกเตอร์ให้เท่ากับ 1 เช่นเดียวกับช่วงเวลา")
body(doc, "หมวดหมู่จะสร้างคะแนนร่วมเมื่อผู้ใช้และทริปมีค่าอยู่ในตำแหน่งเดียวกัน เช่น ผู้ใช้เลือกทะเลและทริปเป็นทะเล ส่วนผู้ใช้เลือกทะเลแต่ทริปเป็นภูเขาจะไม่มีจุดร่วมในส่วนนี้ การกำหนดน้ำหนักหมวดหมู่ไว้สูงสุดจึงช่วยป้องกันไม่ให้ทริปราคาถูกแต่ผิดประเภทแสดงเหนือทริปประเภทที่ผู้ใช้ต้องการ")

heading(doc, "3.5.2 การวิเคราะห์จำนวนมิติของเวกเตอร์")
body(doc, "เวกเตอร์สำหรับจับคู่ผู้ใช้กับทริปประกอบด้วยทั้งหมด 21 มิติ มิใช่ 15 มิติ เนื่องจากงบประมาณและระดับกิจกรรมถูกสร้างเป็นเวกเตอร์อย่างละ 2 มิติ ช่วงเวลามี 4 มิติ และหมวดหมู่มี 13 มิติ")
table(doc, ["กลุ่มคุณลักษณะ", "จำนวนมิติ"], [("งบประมาณ", "2"), ("ระดับกิจกรรม", "2"), ("ช่วงเวลา", "4"), ("หมวดหมู่", "13"), ("รวม", "21")], [10.5, 5.0])
formula(doc, "Vector = [Budget₁…₂, Activity₁…₂, Time₁…₄, Category₁…₁₃]")
body(doc, "โครงสร้าง 21 มิตินี้ใช้กับการจับคู่ผู้ใช้กับทริป ส่วนระบบค้นหาเพื่อนร่วมทางระหว่างผู้ใช้กับผู้ใช้ในโค้ดปัจจุบันใช้วิธีคำนวณความเข้ากันได้อีกกระบวนการหนึ่ง จึงไม่ควรกล่าวว่าใช้ Cosine Similarity เช่นเดียวกัน")

heading(doc, "3.5.3 การกำหนดน้ำหนัก")
body(doc, "ก่อนต่อเวกเตอร์แต่ละกลุ่มเข้าด้วยกัน ระบบคูณสมาชิกในกลุ่มด้วยรากที่สองของน้ำหนัก การใช้ √w มีเหตุผลเพราะเมื่อคำนวณ Dot Product ค่าจากฝั่งผู้ใช้และฝั่งทริปจะถูกคูณกัน ทำให้ √w × √w เท่ากับ w และน้ำหนักที่มีผลจริงยังคงเป็น 35%, 30%, 20% และ 15%")
formula(doc, "V′category = √0.35 Vcategory   |   V′budget = √0.30 Vbudget")
formula(doc, "V′activity = √0.20 Vactivity   |   V′time = √0.15 Vtime")
formula(doc, "Score ≈ 0.35C + 0.30B + 0.20A + 0.15T")

heading(doc, "3.5.4 ขั้นตอนการคำนวณ Cosine Similarity")
heading(doc, "3.5.4.1 ผลคูณจุด (Dot Product)", 2)
body(doc, "Dot Product คือการนำค่าตำแหน่งเดียวกันของเวกเตอร์ผู้ใช้และเวกเตอร์ทริปมาคูณ แล้วรวมผลทั้งหมด ค่าจะสูงเมื่อทั้งสองฝ่ายมีคุณลักษณะตรงกันหลายตำแหน่ง แต่ Dot Product เพียงอย่างเดียวยังไม่ยุติธรรมต่อเวกเตอร์ที่มีขนาดต่างกัน จึงต้องหารด้วยขนาดของเวกเตอร์")
formula(doc, "A · B = Σᵢ₌₁ⁿ AᵢBᵢ")
heading(doc, "3.5.4.2 ขนาดของเวกเตอร์ (Magnitude)", 2)
body(doc, "Magnitude คือความยาวรวมของเวกเตอร์ คำนวณจากรากที่สองของผลรวมกำลังสองของสมาชิกทุกตำแหน่ง ทำหน้าที่ปรับคะแนนให้อยู่บนมาตรฐานเดียวกัน และลดความได้เปรียบของผู้ที่เลือกตัวเลือกจำนวนมาก")
formula(doc, "||A|| = √(Σᵢ₌₁ⁿ Aᵢ²)     และ     ||B|| = √(Σᵢ₌₁ⁿ Bᵢ²)")
heading(doc, "3.5.4.3 ค่าความคล้ายคลึงเชิงโคไซน์", 2)
body(doc, "ระบบนำ Dot Product หารด้วยผลคูณของ Magnitude ทั้งสอง ค่าที่ใช้งานอยู่ระหว่าง 0–1 เพราะคุณลักษณะทั้งหมดไม่ติดลบ ค่าใกล้ 1 หมายถึงเข้ากันได้มาก ค่าใกล้ 0 หมายถึงมีคุณลักษณะร่วมกันน้อย จากนั้นคูณ 100 และปัดเป็นจำนวนเต็มเพื่อแสดง Match Score")
formula(doc, "Cosine Similarity(A,B) = (A · B) ÷ (||A|| × ||B||)")
formula(doc, "Match Score (%) = round(Cosine Similarity × 100)")

heading(doc, "3.5.5 ตัวอย่างการคำนวณ")
body(doc, "ตัวอย่างผู้ใช้มีงบประมาณ 3,000 บาท เลือกระดับกิจกรรม 3–4 รายการต่อวัน ชอบช่วงเช้า เย็น และกลางคืน และสนใจทะเล ส่วนทริปมีงบประมาณ 4,500 บาท ระดับกิจกรรมเท่ากัน มีกิจกรรมช่วงเย็นและกลางคืน และเป็นทริปทะเล ระบบได้ความคล้ายคลึงรายด้านโดยประมาณดังนี้")
table(doc, ["ด้านที่เปรียบเทียบ", "ความคล้ายคลึง", "ส่วนคะแนนตามน้ำหนัก"], [
    ("หมวดหมู่", "1.0000 หรือ 100%", "1.0000 × 0.35 = 0.3500"),
    ("งบประมาณ", "0.9948 หรือ 99%", "0.9948 × 0.30 = 0.2984"),
    ("ระดับกิจกรรม", "1.0000 หรือ 100%", "1.0000 × 0.20 = 0.2000"),
    ("ช่วงเวลา", "0.8165 หรือ 82%", "0.8165 × 0.15 = 0.1225"),
], [4.4, 4.6, 6.5])
formula(doc, "รวม = 0.3500 + 0.2984 + 0.2000 + 0.1225 = 0.9709 ≈ 97%")
body(doc, "คะแนน 97% ไม่ได้หมายความว่าผู้ใช้เหมือนกับทริปทุกประการ แต่หมายถึงเมื่อประเมินทั้ง 4 ด้านตามน้ำหนักที่กำหนดแล้ว ผู้ใช้มีความเหมาะสมกับทริปนี้ในระดับ 97%")

heading(doc, "3.5.6 เหตุผลที่เลือกใช้ Cosine Similarity")
body(doc, "Cosine Similarity เหมาะกับระบบนี้เพราะสามารถเปรียบเทียบคุณลักษณะหลายด้านพร้อมกัน โดยเน้นทิศทางของรูปแบบความชอบมากกว่าขนาดของตัวเลข การปรับขนาดเวกเตอร์ช่วยลดความได้เปรียบของผู้ที่เลือกคำตอบจำนวนมาก อีกทั้งสามารถแยกอธิบายที่มาของคะแนนตามหมวดหมู่ งบประมาณ กิจกรรม และช่วงเวลาได้ ทำให้ตรวจสอบกลไกของระบบได้ง่าย")
body(doc, "นอกจากนี้ สูตรมีความซับซ้อนไม่สูง เหมาะกับการคำนวณทริปหลายรายการ และสามารถขยายเวกเตอร์เพื่อรองรับคุณลักษณะใหม่ในอนาคตได้ อย่างไรก็ตาม Cosine Similarity เป็นเพียงกลไกวัดความคล้ายคลึง มิใช่ปัญญาประดิษฐ์ที่เรียนรู้น้ำหนักด้วยตนเอง")

heading(doc, "3.5.7 เงื่อนไขทางธุรกิจหลังการคำนวณ")
body(doc, "หลังได้คะแนน Cosine Similarity ระบบตรวจสอบเงื่อนไขการเข้าร่วมจริงเพิ่มเติม เงื่อนไขเหล่านี้เป็นกฎหลังการคำนวณ ไม่ใช่ส่วนหนึ่งของสูตร Cosine")
bullet(doc, "หากจำนวนสมาชิกเต็ม ระบบกำหนดคะแนนที่แสดงเป็น 0")
bullet(doc, "หากทริปสิ้นสุดหรือพ้นวันเดินทางแล้ว ระบบกำหนดคะแนนที่แสดงเป็น 0")
bullet(doc, "หากงบทริปสูงกว่างบผู้ใช้เกิน 2 เท่า ระบบจำกัดคะแนนรวมไม่ให้เกิน 39%")
body(doc, "การจำกัดคะแนนช่วยป้องกันไม่ให้ทริปที่ผู้ใช้อาจรับภาระค่าใช้จ่ายไม่ได้แสดงคะแนนสูงเพียงเพราะหมวดหมู่และกิจกรรมตรงกัน จึงควรอธิบายแยกจากสูตร Cosine เพื่อไม่ให้เข้าใจว่าค่าความคล้ายคลึงถูกคำนวณผิด")

heading(doc, "3.5.8 ขอบเขตและข้อจำกัด")
body(doc, "คะแนนเป็นการประมาณความเหมาะสมจากข้อมูลที่ผู้ใช้และผู้สร้างทริประบุ ไม่สามารถรับประกันความพึงพอใจจริงทั้งหมด เพราะยังมีปัจจัยอื่น เช่น บุคลิกภาพ ความปลอดภัย ประเภทที่พัก วิธีเดินทาง และพฤติกรรมสมาชิกที่ยังไม่ได้รวมในเวกเตอร์")
body(doc, "น้ำหนัก 35%, 30%, 20% และ 15% เป็นค่าที่ออกแบบตามเหตุผลของระบบ ยังไม่ใช่ผลจากการเรียนรู้ด้วยชุดข้อมูลขนาดใหญ่ หากต้องการยืนยันว่าชุดน้ำหนักใดดีที่สุด ควรเก็บข้อมูลการกดดู การเข้าร่วม การยกเลิก และความพึงพอใจหลังเดินทาง แล้วเปรียบเทียบชุดน้ำหนักหลายแบบด้วยตัวชี้วัดที่เหมาะสม")

doc.save(OUT)
print(OUT)
