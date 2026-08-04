from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/Users/worakanp/Desktop/go-with-us-1')
IMG = ROOT / 'reports/ui-screenshots'
OUT = ROOT / 'reports/GoWithUs_UX_UI_Report.docx'

NAVY = '203F73'
BLUE = '5B8DEF'
INK = '243044'
MUTED = '667085'
PALE = 'F2F4F7'
MINT = 'DFF3EA'
RED = 'B42318'
FONT = 'Arial Unicode MS'


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:cs'), FONT)
    run.font.complex_script = True
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'th-TH')
    lang.set(qn('w:eastAsia'), 'th-TH')
    lang.set(qn('w:bidi'), 'th-TH')
    run._element.get_or_add_rPr().append(lang)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc.append(tc_mar)
    for tag, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value)); node.set(qn('w:type'), 'dxa')


def p(doc, text='', size=11, bold=False, color=INK, after=6, before=0,
      align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, keep=False):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.10
    para.paragraph_format.keep_with_next = keep
    set_font(para.add_run(text), size, bold, color, italic)
    return para


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 12}
    before = {1: 16, 2: 12, 3: 8}
    after = {1: 8, 2: 6, 3: 4}
    return p(doc, text, sizes[level], True, NAVY, after[level], before[level], keep=True)


def bullet(doc, text, color=INK):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(.5)
    para.paragraph_format.first_line_indent = Inches(-.25)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.167
    set_font(para.add_run(text), 10.5, False, color)


def callout(doc, title, text, fill=MINT):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.5)
    c = t.cell(0, 0); shade(c, fill); margins(c, 130, 180, 130, 180)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    q = c.paragraphs[0]; q.paragraph_format.space_after = Pt(3)
    set_font(q.add_run(title), 10.5, True, NAVY)
    q2 = c.add_paragraph(); q2.paragraph_format.space_after = Pt(0); q2.paragraph_format.line_spacing = 1.10
    set_font(q2.add_run(text), 10, False, INK)


def add_img(doc, name, caption, width=2.6):
    path = IMG / name
    if not path.exists():
        return
    q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q.paragraph_format.space_before = Pt(4); q.paragraph_format.space_after = Pt(3)
    q.add_run().add_picture(str(path), width=Inches(width))
    p(doc, caption, 9, False, MUTED, 8, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)


def image_pair(doc, items):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, (name, caption) in enumerate(items):
        c = t.cell(0, i); c.width = Inches(3.2); margins(c, 40, 60, 40, 60)
        q = c.paragraphs[0]; q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.add_run().add_picture(str(IMG / name), width=Inches(2.15))
        cap = c.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cap.add_run(caption), 8.5, False, MUTED, True)


def page(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1); sec.right_margin = Inches(1)
sec.header_distance = Inches(.492); sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles['Normal']; normal.font.name = FONT; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for style_name in ['List Bullet', 'List Number']:
    s = styles[style_name]; s.font.name = FONT; s.font.size = Pt(10.5)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run('GoWithUs · UX/UI Review'), 8.5, False, MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run('รายงานจากการทดลองบน iPhone 13 Simulator และการตรวจโค้ด SwiftUI · 3 สิงหาคม 2026'), 8, False, MUTED)

# Editorial cover
p(doc, 'PRODUCT EXPERIENCE REPORT', 10, True, BLUE, 14, 45, WD_ALIGN_PARAGRAPH.CENTER)
p(doc, 'GoWithUs', 30, True, NAVY, 5, align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, 'รายงานการตรวจ Interface และ UX/UI แบบหน้า‑ต่อ‑หน้า', 16, False, INK, 10, align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, 'วิเคราะห์จากการใช้งานจริงบน Simulator และอ่านโค้ด iOS (SwiftUI)', 10.5, False, MUTED, 18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_img(doc, '01-current.png', 'หน้าหลักของแอพขณะทดสอบ', 2.1)
p(doc, 'เวอร์ชันตรวจสอบ: 3 สิงหาคม 2026', 9.5, True, MUTED, 0, 18, WD_ALIGN_PARAGRAPH.CENTER)

page(doc)
heading(doc, '1. บทสรุปสำหรับเจ้าของแอพ')
p(doc, 'GoWithUs มีโครงสร้างผลิตภัณฑ์ชัด: ค้นหาทริป → ดูรายละเอียด → แสดงความสนใจ/เข้าร่วม → หาเพื่อน → สนทนา และสร้างทริปเองได้ จุดแข็งคือฟังก์ชันครบและภาษาไทยเป็นมิตร แต่ประสบการณ์ยังดูเหมือนหลายชุดดีไซน์ถูกนำมารวมกัน จึงเกิดความไม่สม่ำเสมอของสี ระยะห่าง ปุ่ม และรูปแบบนำทาง')
callout(doc, 'ข้อสรุปหลัก', 'ควรยึด Navy #203F73 เป็นสีหลัก, Mist Blue #5B8DEF เป็นสีโต้ตอบ, พื้นหลัง #F7F8FA และใช้เขียวเฉพาะสถานะสำเร็จ/เปอร์เซ็นต์แมตช์ เพื่อลดความแสบตาและทำให้ภาพรวมเป็นมินิมอลตามทิศทางที่เลือกไว้')
heading(doc, 'สิ่งที่ควรแก้ก่อน', 2)
bullet(doc, 'P0 — แก้ Full-screen Image Viewer ที่กดปิดแล้วอาจค้างเป็นพื้นดำ/รูปเดิม เพราะเป็นปัญหาที่ตัดการใช้งาน')
bullet(doc, 'P1 — รวม Navigation ให้เป็นภาษาดีไซน์เดียวกัน: บางหน้าใช้หัวแอพเอง บางหน้าใช้ NavigationStack และบางหน้ามี “Back” ซ้ำกับ “ยกเลิก”')
bullet(doc, 'P1 — ทำระบบสีและ component กลางสำหรับปุ่ม, chip, card, empty state และสถานะ เพื่อให้ทุกหน้าเชื่อมโยงกัน')
bullet(doc, 'P1 — ทำ action ของหน้า Find Buddy ให้เห็นชัด เพิ่มคำแนะนำ “ปัดซ้าย/ขวา” และปุ่ม Like/Pass ในระดับสายตา')
bullet(doc, 'P2 — ลดพื้นที่ว่างผิดสัดส่วนในหน้าโปรไฟล์และ AI Chat พร้อมแสดง next action ที่ชัดเจน')
heading(doc, 'ขอบเขตและวิธีตรวจ', 2)
p(doc, 'ทดสอบบน iPhone 13 Simulator (iOS 18.0), เปิดและเลื่อนหน้าที่เข้าถึงได้จริง พร้อมอ่าน View, ViewModel และ Service ที่เกี่ยวข้อง ภาพในรายงานเป็นหลักฐานจาก Simulator; หน้าที่เปิดด้วยการควบคุมภายนอกไม่ได้จะระบุว่าอ้างอิงจากโค้ด')

page(doc)
heading(doc, '2. โครงสร้างการนำทางและ Interface')
p(doc, 'โค้ด ContentView.swift ใช้หน้าหลักแบบ persistent tabs แล้วซ่อน/เปิดด้วย opacity และ hit testing มี bottom tab 5 จุด: หน้าแรก, แมตช์ทริป, สร้าง, แชท, โปรไฟล์ และมี Side Menu สำหรับรายการโปรด หาเพื่อน ทริปของฉัน AI และการตั้งค่า')
image_pair(doc, [('01-current.png','หน้าหลัก + Bottom Tab'), ('06-side-menu.png','Side Menu')])
heading(doc, 'ข้อดี', 2)
bullet(doc, 'ฟังก์ชันหลักเข้าถึงได้ทั้งแถบล่างและเมนูด้านข้าง เหมาะกับแอพที่มีหลายเส้นทางใช้งาน')
bullet(doc, 'ปุ่มสร้างทริปตรงกลางเด่นและสื่อความหมายทันที')
heading(doc, 'สิ่งที่ควรปรับ', 2)
bullet(doc, 'เมนูซ้ำกันหลายตำแหน่งเพิ่มภาระการเรียนรู้ ควรกำหนดว่า Bottom Tab คือ 4–5 งานหลัก ส่วน Side Menu เก็บเฉพาะงานรอง')
bullet(doc, 'แถบล่างลอยทับเนื้อหาในบางจังหวะ ควรเพิ่ม safe-area inset ที่คำนวณจากความสูงจริงของ tab bar')
bullet(doc, 'ใช้ชื่อและไอคอนให้คงที่ เช่น “แมตช์ทริป” ต้องหมายถึงทริปแนะนำ ไม่ปะปนกับ “หาเพื่อน”')

page(doc)
heading(doc, '3. หน้าแรก — ค้นหาและเลือกทริป')
add_img(doc, '01-current.png', 'ภาพทดลองจริง: HomeView', 2.75)
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'HomeView แสดงโลโก้ เมนู รูปแบบ grid การแจ้งเตือน ฟิลเตอร์จังหวัด/วันเดินทาง/สไตล์ ช่องค้นหา segmented control (แนะนำ/มาใหม่/ยอดนิยม) และ TripCard หลายใบ การแตะการ์ดนำไป TripDetailView')
heading(doc, 'UX/UI ที่พบ', 2)
bullet(doc, 'ลำดับข้อมูลดีและค้นหาได้หลายทาง แต่บริเวณบนมี control จำนวนมาก ทำให้ภาพรวมแน่น')
bullet(doc, 'แท็กบนรูปหลายชิ้นและเปอร์เซ็นต์ 0% แย่งความสนใจจากชื่อทริป ควรซ่อนค่า 0% หรืออธิบายสาเหตุ')
bullet(doc, 'การ์ดใบที่สองถูกแถบล่างทับ ควรเพิ่ม bottom content inset')
heading(doc, 'ข้อเสนอ', 2)
p(doc, 'ลดฟิลเตอร์บนหน้าให้เหลือ 2 ปุ่มหลักและปุ่ม “ตัวกรองทั้งหมด”; จำกัด tag บนภาพ 2 รายการ + “+N”; ใช้ภาพ hero ที่มืดลงเล็กน้อยเฉพาะบริเวณข้อความเพื่อเพิ่ม contrast')

page(doc)
heading(doc, '4. รายละเอียดทริปและ Gallery')
image_pair(doc, [('03-profile-test.png','Trip Detail'), ('04-chat.png','Full-screen Gallery')])
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'TripDetailView แบ่งเนื้อหาเป็น hero image, ชื่อ/สถานที่, badge, ความเข้ากันได้, งบประมาณ จำนวนคน ระยะเวลา ผู้จัด รายละเอียด tag itinerary gallery ผู้เข้าร่วม และ action bar สำหรับ “สนใจ/จะไปด้วย/ออกจากทริป”')
heading(doc, 'ปัญหาสำคัญที่พบ', 2)
bullet(doc, 'หลังเปิดรูปเต็มจอและกดปิด มีกรณีพื้นดำหรือรูปเดิมค้างอยู่ (ภาพถัดจากการกดปิดยังไม่กลับหน้ารายละเอียด) ควรตรวจ state ของ fullScreenCover/item และ dismiss')
bullet(doc, 'หน้าเนื้อหายาวมาก ควรมี sticky action bar และ anchor/section spacing ที่สม่ำเสมอ')
bullet(doc, 'คำว่า “งบประมาณ” ควรแสดงชนิดให้ชัดเสมอว่า “ต่อคน” หรือ “รวมทั้งทริป”; การคำนวณประมาณ ×3 ควรมีคำอธิบายที่มา ไม่ให้ผู้ใช้คิดว่าเป็นค่าจริง')

page(doc)
heading(doc, '5. หาเพื่อน — Buddy Matching')
add_img(doc, '07-find-buddy.png', 'ภาพทดลองจริง: Find Buddy แสดงการ์ดผู้ใช้และคะแนน Match', 2.8)
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'FindBuddyView โหลดผู้ใช้จาก FindBuddyViewModel/MatchService แสดงการ์ดแบบ swipe มีคะแนนความเข้ากันได้ ความสนใจ ปัดเพื่อ Like/Nope และเปิดรายการ Mutual Matches เพื่อเริ่มแชทเมื่อแมตช์กัน')
heading(doc, 'UX/UI ที่พบ', 2)
bullet(doc, 'รูปบุคคลและคะแนนเด่นดี แต่การ์ดกินพื้นที่เกือบทั้งหน้าจอจน action อยู่ต่ำและมองไม่เห็นทันที')
bullet(doc, 'พฤติกรรมปัดยังไม่ discoverable สำหรับผู้ใช้ใหม่')
heading(doc, 'ข้อเสนอ', 2)
p(doc, 'ลดความสูงการ์ด 10–15%, ตรึงปุ่ม Pass / Like ด้านล่าง, ใส่คำแนะนำครั้งแรก และใช้สีแดง/เขียวเฉพาะ action ไม่ใช้เป็นสีประจำหน้าทั้งหมด')

page(doc)
heading(doc, '6. ทริปของฉันและ Empty State')
image_pair(doc, [('08-my-trips.png','สร้างเอง'), ('09-my-trips-joined.png','เข้าร่วมแล้ว')])
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'MyTripsView มี segmented picker 3 หมวด: สร้างเอง, เข้าร่วมแล้ว, รายการโปรด โดยโหลดข้อมูลผ่าน MyTripsViewModel และมี CTA สร้างทริปในหมวดแรก')
heading(doc, 'UX/UI ที่พบ', 2)
bullet(doc, 'ข้อความ empty state ตรงไปตรงมาและ CTA สร้างทริปเหมาะกับบริบท')
bullet(doc, 'พื้นที่ว่างมากและแต่ละหมวดแทบเหมือนกัน ควรมี illustration ขนาดเล็กและข้อความแนะนำเฉพาะหมวด')
bullet(doc, '“รายการโปรด” มีทั้งใน Side Menu และ tab นี้ ควรเลือกจุดหลักหนึ่งแห่งเพื่อลดความซ้ำซ้อน')

page(doc)
heading(doc, '7. สร้างทริป')
image_pair(doc, [('10-create-trip.png','ส่วนต้นของฟอร์ม'), ('11-create-trip-scrolled.png','งบและรูปแบบทริป')])
add_img(doc, '12-create-trip-bottom.png', 'ส่วนท้าย: ช่วงเวลา แผนการเดินทาง และปุ่มสร้างทริป', 2.15)
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'CreateTripView รองรับหลายรูป จังหวัด สไตล์ รายละเอียด AI ช่วยจัดทริป tag ความเป็นสาธารณะ วันเริ่ม/สิ้นสุด งบประมาณแบบต่อคนหรือรวม จำนวนคน activity style ช่วงเวลา itinerary และการแก้ไข/ลบทริป')
heading(doc, 'UX/UI ที่พบและข้อเสนอ', 2)
bullet(doc, 'ฟังก์ชันครบมาก แต่เป็นฟอร์มยาว ควรแบ่งเป็น 3 ขั้น: ข้อมูลหลัก → งบ/ผู้ร่วมทริป → แผนและตรวจสอบ')
bullet(doc, 'ส่วนหัวมี “Back” จากระบบและ “ยกเลิก” ของหน้าเดียวกัน ควรเหลือรูปแบบเดียว')
bullet(doc, 'ประเภทงบ “ต่อคน/ต่อทริป” เป็นจุดสำคัญ ควรแสดงตัวอย่างใต้ช่อง เช่น “ประมาณ 1,500 บาท × 3 คน = 4,500 บาท” และให้ผู้ใช้ยืนยันฐานการคำนวณ')
bullet(doc, 'ปุ่มสร้างทริปควรติดล่างแบบ safe-area และมีสรุปข้อมูลที่ยังไม่ครบก่อนกด')

page(doc)
heading(doc, '8. AI ผู้ช่วยวางแผน')
add_img(doc, '13-ai-chat.png', 'ภาพทดลองจริง: AI Chat', 2.8)
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'AIChatView เก็บประวัติข้อความ ส่ง prompt ผ่าน AIChatViewModel/GeminiService และสามารถแปลงคำตอบเป็น TripDraft เพื่อสร้างทันทีหรือเปิด CreateTripView แก้ไขก่อนบันทึก')
heading(doc, 'UX/UI ที่พบ', 2)
bullet(doc, 'เส้นทาง “คุย → ได้ร่าง → สร้าง/แก้ไข” เป็นจุดขายที่ดีและเชื่อมกับฟอร์มสร้างทริปได้')
bullet(doc, 'หน้าว่างมากก่อนเริ่มสนทนา ควรเพิ่ม suggestion chips เช่น “เชียงใหม่ 3 วัน งบ 5,000” และอธิบายว่า AI จะถามอะไรบ้าง')
bullet(doc, 'ต้องแยกข้อมูล “AI ประมาณให้” ออกจากค่าที่ผู้ใช้ยืนยันแล้ว และให้แก้ทุกช่องก่อนสร้างจริง')

page(doc)
heading(doc, '9. แชท')
add_img(doc, '18-chat-list.png', 'ภาพทดลองจริง: รายการสนทนาที่ยังว่าง', 2.8)
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'ChatView แสดง Mutual Matches ด้านบนและรายการข้อความล่าสุดด้านล่าง ใช้ ChatViewModel/MessageService พร้อม unread count; แตะผู้ใช้เพื่อเปิด ChatDetailView')
heading(doc, 'UX/UI ที่พบและข้อเสนอ', 2)
bullet(doc, 'Empty state เข้าใจง่าย แต่ไม่มี CTA ไปหาเพื่อนหรือค้นหาทริป ควรเพิ่มปุ่ม “หาเพื่อนใหม่”')
bullet(doc, 'เมื่อมีข้อมูล ควรแบ่งหัวข้อ “เพื่อนใหม่ที่แมตช์กัน” กับ “ข้อความล่าสุด” อย่างชัด และไม่แสดงส่วนว่างที่ไม่มีรายการ')
bullet(doc, 'สถานะ unread ควรใช้สี accent เดียวกับระบบ ไม่ใช้สีใหม่เพิ่ม')

page(doc)
heading(doc, '10. โปรไฟล์และแก้ไขโปรไฟล์')
image_pair(doc, [('14-profile.png','หน้าโปรไฟล์'), ('16-edit-profile.png','แก้ไขข้อมูลส่วนตัว')])
add_img(doc, '17-edit-profile-bottom.png', 'สไตล์การเที่ยวและความเป็นส่วนตัว', 2.1)
heading(doc, 'Interface และการทำงาน', 2)
p(doc, 'ProfileView มี header, สถานะยืนยันตัวตน, ข้อมูลผู้ใช้ และทริปของผู้ใช้; EditProfileView รองรับรูปสูงสุด 6 รูป username เพศ วันเกิด bio แบบทดสอบไลฟ์สไตล์ สไตล์ท่องเที่ยว และ privacy settings')
heading(doc, 'UX/UI ที่พบ', 2)
bullet(doc, 'หน้าหลักมีพื้นที่ว่างมากผิดสัดส่วนเมื่อข้อมูลน้อย ควรยุบส่วนที่ไม่มีข้อมูลและแสดง completion checklist')
bullet(doc, 'แถบยืนยันตัวตนเปิดเว็บภายนอกผ่าน Link ซึ่งทำงานตามโค้ด แต่ผู้ใช้หลุดบริบทแอพ ควรใช้ in-app browser และบอกก่อนเปิด')
bullet(doc, 'หน้าแก้ไขมีข้อมูลมาก ควรแบ่ง section ให้เห็น hierarchy และมี Save แบบ sticky พร้อมสถานะบันทึกสำเร็จ')

page(doc)
heading(doc, '11. การแจ้งเตือนและหน้าที่ตรวจจากโค้ด')
p(doc, 'NotificationView เปิดเป็น sheet จากหน้า Home มีรายการ NotificationRow, mark as read, ลบรายชิ้น, ลบทั้งหมด และปิด ขณะทดสอบการควบคุม Simulator ภายนอกไม่สามารถเปิด sheet นี้ได้อย่างเสถียร จึงประเมินโครงสร้างจากโค้ด')
heading(doc, 'Match Trip', 2)
p(doc, 'MatchTripView แสดงการ์ดทริปแบบ Tinder swipe พร้อม empty/error/retry state และโหลดคะแนนผ่าน MatchTripViewModel ควรเพิ่มปุ่ม action ที่แตะได้ควบคู่กับ gesture และข้อความอธิบายปัจจัยคะแนน')
heading(doc, 'Authentication / Onboarding', 2)
p(doc, 'LoginView, RegisterView, OTPVerificationView และ OnboardingView มีอยู่ในโค้ด แต่ไม่ได้ออกจากระบบเพื่อเก็บภาพ เพราะจะเปลี่ยน session ผู้ใช้ ควรทดสอบแยกในบัญชีทดสอบ โดยเน้น validation, keyboard, error state, OTP timeout และสิทธิ์ความเป็นส่วนตัว')
heading(doc, 'Settings', 2)
p(doc, 'SideMenuView มีภาษาไทย/อังกฤษและ vibration setting ควรใช้ switch/selection ที่แสดงสถานะชัด และแจ้งผลของการเปลี่ยนภาษาว่ามีผลทันทีหรือหลังเปิดแอพใหม่')

page(doc)
heading(doc, '12. ระบบสีมินิมอลที่แนะนำ')
p(doc, 'เป้าหมายคือดูนุ่ม สงบ และไม่ “ปี้ด” โดยยังคงบุคลิกของ GoWithUs สีเข้มควรใช้เป็นฐาน ส่วนสีสดใช้เฉพาะ interaction และ feedback')
palette = [
    ('Primary Navy', '#203F73', 'หัวข้อ โลโก้ navigation และปุ่มหลัก'),
    ('Mist Blue', '#5B8DEF', 'selection, link, focus และ icon ที่ active'),
    ('Canvas', '#F7F8FA', 'พื้นหลังหลัก'),
    ('Surface', '#FFFFFF', 'การ์ดและ sheet'),
    ('Ink', '#243044', 'ข้อความหลัก'),
    ('Muted', '#667085', 'ข้อความรอง'),
    ('Success', '#2E7D64', 'สำเร็จ/แมตช์เท่านั้น'),
    ('Danger', '#B42318', 'ลบ/ผิดพลาดเท่านั้น'),
]
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
hdr = t.rows[0].cells
for i, txt in enumerate(['สี', 'รหัส', 'การใช้งาน']):
    shade(hdr[i], PALE); margins(hdr[i]); set_font(hdr[i].paragraphs[0].add_run(txt), 10, True, NAVY)
for name, code, use in palette:
    cells = t.add_row().cells
    for c in cells: margins(c)
    shade(cells[0], code[1:])
    set_font(cells[0].paragraphs[0].add_run(name), 9.5, True, 'FFFFFF' if code not in ['#F7F8FA','#FFFFFF'] else INK)
    set_font(cells[1].paragraphs[0].add_run(code), 9.5, False, INK)
    set_font(cells[2].paragraphs[0].add_run(use), 9.5, False, INK)
heading(doc, 'กติกาการใช้สี', 2)
bullet(doc, 'หนึ่งหน้าควรมี accent หลักไม่เกิน 1 สี; สีเขียว/แดงเป็น semantic state เท่านั้น')
bullet(doc, 'พื้นหลังใช้ Canvas และวางข้อมูลบน Surface พร้อมเส้นขอบ #E4E7EC แทนเงาหนัก')
bullet(doc, 'ตัวอักษรหลักใช้ Ink ไม่ใช้ดำสนิท; ข้อความรองใช้ Muted และรักษา contrast ให้อ่านได้')

page(doc)
heading(doc, '13. Design System และลำดับการปรับปรุง')
heading(doc, 'Component ที่ควรทำเป็นชุดกลาง', 2)
bullet(doc, 'GWPrimaryButton / GWSecondaryButton / GWDestructiveButton — สูง 48–52 pt, radius 14–16 pt')
bullet(doc, 'GWCard — พื้นขาว ขอบบาง radius 18–20 pt และ padding 16 pt')
bullet(doc, 'GWChip — selected/unselected, จำกัดสีและความสูงให้เท่ากันทุกหน้า')
bullet(doc, 'GWEmptyState — icon, title, description และ CTA 1 ปุ่ม')
bullet(doc, 'GWNavigationHeader — กติกาปุ่ม back, title และ trailing action แบบเดียวกัน')
bullet(doc, 'Spacing tokens: 4, 8, 12, 16, 24, 32 pt; typography tokens: 12, 14, 16, 20, 28 pt')
heading(doc, 'Roadmap ที่แนะนำ', 2)
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
for i, txt in enumerate(['ระยะ', 'งาน', 'ผลลัพธ์']):
    shade(t.rows[0].cells[i], PALE); margins(t.rows[0].cells[i]); set_font(t.rows[0].cells[i].paragraphs[0].add_run(txt), 10, True, NAVY)
rows = [
    ('สัปดาห์ 1', 'แก้ Gallery ค้าง, safe area, navigation ซ้ำ', 'ลด bug และทางตัน'),
    ('สัปดาห์ 2', 'ลง color/spacing/type tokens และ component หลัก', 'สีและรูปแบบสอดคล้องกัน'),
    ('สัปดาห์ 3', 'ปรับ Home, Create Trip, Find Buddy', 'งานหลักใช้ง่ายขึ้น'),
    ('สัปดาห์ 4', 'ทดสอบ auth, accessibility, error/empty/loading', 'พร้อมทดสอบผู้ใช้จริง'),
]
for row in rows:
    cells = t.add_row().cells
    for i, txt in enumerate(row): margins(cells[i]); set_font(cells[i].paragraphs[0].add_run(txt), 9.5, i == 0, INK)
heading(doc, 'เกณฑ์สำเร็จ', 2)
bullet(doc, 'ผู้ใช้ใหม่สร้างทริปสำเร็จโดยไม่ถามผู้ทดสอบ และเข้าใจชนิดงบประมาณก่อนบันทึก')
bullet(doc, 'ผู้ใช้หาเพื่อนและเริ่มแชทได้ภายใน 3 interaction หลังเข้าเมนู')
bullet(doc, 'ทุกหน้ามี loading, empty, error และ success state ที่ใช้ภาษาและสีชุดเดียวกัน')
callout(doc, 'บทสรุป', 'ฐานฟังก์ชันของ GoWithUs ดีและมีเอกลักษณ์ชัด งานถัดไปไม่จำเป็นต้องเพิ่มสีหรือ effect แต่ควรลดจำนวนรูปแบบ ทำ hierarchy ให้แน่น และแก้ทางตันก่อน การปรับระบบกลางรอบเดียวจะทำให้ทุกหน้าดูเป็นแอพเดียวกันทันที')

doc.core_properties.title = 'GoWithUs UX/UI Report'
doc.core_properties.subject = 'Simulator review and SwiftUI code analysis'
doc.core_properties.author = 'Codex'
doc.save(OUT)
print(OUT)
