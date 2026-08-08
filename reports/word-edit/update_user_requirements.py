from docx import Document


SOURCE = "reports/word-edit/เว็บแอพแก้ล่าสุด_ฉบับแก้ไข.docx"


def replace_in_runs(paragraph, old, new):
    """Replace text without resetting the paragraph's existing formatting."""
    full_text = "".join(run.text for run in paragraph.runs)
    if old not in full_text:
        raise ValueError(f"Text not found: {old}")

    # The two target phrases each occupy a single run in this document, so
    # editing that run preserves the surrounding tabs, page number, and style.
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return

    raise ValueError(f"Target spans runs unexpectedly: {old}")


document = Document(SOURCE)
old_heading = "การศึกษาปัญหาและกำหนดกรอบแนวคิดของระบบ"
new_heading = "การวิเคราะห์ความต้องการของระบบ"

# Heading in Chapter 3.
replace_in_runs(document.paragraphs[228], old_heading, new_heading)

# Matching entry in the manually maintained table of contents.
replace_in_runs(document.paragraphs[60], old_heading, new_heading)

document.save(SOURCE)
